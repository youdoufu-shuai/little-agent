import os
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

class DocumentGenerator:
    """
    处理混合内容的 .docx 和 .pdf 文档生成。
    """
    
    @staticmethod
    def _register_chinese_font():
        """注册中文字体以支持 PDF 中文输出"""
        font_name = "STHeiti" # 默认回退名称
        available_fonts = [
            ("/System/Library/Fonts/STHeiti Light.ttc", "STHeiti"),
            ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
            ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
            ("/System/Library/Fonts/Supplemental/Arial Unicode.ttf", "ArialUnicode"),
        ]
        
        registered = False
        for path, name in available_fonts:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(name, path))
                    font_name = name
                    registered = True
                    break
                except Exception:
                    continue
        
        return font_name

    @staticmethod
    def generate(filename: str, file_type: str, content: List[Dict[str, Any]], style_config: Optional[Dict[str, Any]] = None) -> str:
        """
        生成文档。
        
        Args:
            filename: 输出文件名（基本名称）。
            file_type: 'docx' 或 'pdf'。
            content: 内容块列表。
            style_config: 全局样式设置。
            
        Returns:
            生成文件的绝对路径。
        """
        # 确保输出目录存在
        # 使用相对于当前文件 (core/doc_generator.py) 的路径来定位 web/files
        current_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(current_dir)
        output_dir = os.path.join(project_root, "web", "files")
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # 确保扩展名匹配
        if not filename.lower().endswith(f".{file_type}"):
            filename += f".{file_type}"
            
        file_path = os.path.join(output_dir, filename)
        
        # 跟踪下载的临时文件以便清理
        downloaded_files = []
        
        try:
            if file_type.lower() == 'docx':
                result = DocumentGenerator._generate_docx(file_path, content, style_config, downloaded_files)
            elif file_type.lower() == 'pdf':
                result = DocumentGenerator._generate_pdf(file_path, content, style_config, downloaded_files)
            else:
                raise ValueError("Unsupported file type. Use 'docx' or 'pdf'.")
                
            return result
        finally:
            # 清理下载的临时图片
            print(f"Cleaning up {len(downloaded_files)} temporary images...")
            for temp_file in downloaded_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        print(f"Deleted temporary file: {temp_file}")
                except Exception as e:
                    print(f"Failed to delete temporary file {temp_file}: {e}")

    @staticmethod
    def _resolve_image_path(path: str, downloaded_files: Optional[List[str]] = None) -> str:
        """
        解析图像路径。
        1. 如果是 URL，下载到 web/images/web_crawled (如果不存在)。
        2. 如果是本地路径，尝试解析为绝对路径。
        """
        import requests
        from urllib.parse import urlparse
        import hashlib

        # 0. 处理 /static/ 路径 (映射到本地 web 目录)
        if path.startswith("/static/"):
            # 获取 web 目录的绝对路径
            # core/doc_generator.py -> core/ -> project_root -> web/
            current_dir = os.path.dirname(os.path.abspath(__file__))
            project_root = os.path.dirname(current_dir)
            web_dir = os.path.join(project_root, "web")
            
            # 移除 /static/ 前缀 (8个字符) 并拼接
            # /static/images/foo.png -> images/foo.png
            relative_path = path[8:] 
            local_path = os.path.join(web_dir, relative_path)
            return local_path

        # 检查是否为 URL
        if path.startswith("http://") or path.startswith("https://"):
            try:
                # 确定保存路径
                # 使用 URL 哈希作为文件名以避免冲突
                parsed_url = urlparse(path)
                filename = os.path.basename(parsed_url.path)
                if not filename or "." not in filename:
                    # 如果 URL 中没有明确的文件名，使用 hash
                    ext = ".jpg" # 默认
                    filename = hashlib.md5(path.encode('utf-8')).hexdigest() + ext
                else:
                    # 清理文件名
                    filename = "".join(c for c in filename if c.isalnum() or c in "._-")
                
                # 保存到 web/images/web_crawled
                save_dir = os.path.join(os.getcwd(), "web", "images", "web_crawled")
                if not os.path.exists(save_dir):
                    os.makedirs(save_dir)
                
                local_path = os.path.join(save_dir, filename)
                
                # 如果文件已存在，直接使用（简单的缓存机制）
                if os.path.exists(local_path):
                    if downloaded_files is not None and "web_crawled" in local_path:
                        downloaded_files.append(local_path)
                    return local_path
                
                # 下载
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(path, headers=headers, timeout=30)
                
                if response.status_code == 200:
                    content_type = response.headers.get('Content-Type', '').lower()
                    
                    # 如果是 HTML 页面，尝试提取 og:image
                    if 'text/html' in content_type:
                        print(f"URL is HTML page, attempting to extract og:image from {path}")
                        import re
                        html_content = response.text
                        # 查找 <meta property="og:image" content="..."> 或 <meta name="twitter:image" content="...">
                        og_image_match = re.search(r'<meta\s+(?:property|name)=["\'](?:og:image|twitter:image)["\']\s+content=["\']([^"\']+)["\']', html_content, re.IGNORECASE)
                        
                        if og_image_match:
                            image_url = og_image_match.group(1)
                            # 处理相对路径
                            if not image_url.startswith('http'):
                                from urllib.parse import urljoin
                                image_url = urljoin(path, image_url)
                            
                            print(f"Found image URL in meta tags: {image_url}")
                            # 递归调用自身下载提取的图片 URL
                            # 但为了避免无限递归（如果那个 URL 又是 HTML），我们直接下载它
                            img_response = requests.get(image_url, headers=headers, timeout=30)
                            if img_response.status_code == 200 and 'image' in img_response.headers.get('Content-Type', '').lower():
                                with open(local_path, "wb") as f:
                                    f.write(img_response.content)
                                if downloaded_files is not None and "web_crawled" in local_path:
                                    downloaded_files.append(local_path)
                                return local_path
                            else:
                                print(f"Failed to download extracted image {image_url}")
                        else:
                            print("No og:image found in HTML")
                            
                    # 如果是图片直接保存
                    else:
                        with open(local_path, "wb") as f:
                            f.write(response.content)
                        if downloaded_files is not None and "web_crawled" in local_path:
                            downloaded_files.append(local_path)
                        return local_path
                        
                else:
                    print(f"Failed to download {path}: Status {response.status_code}")
                    return path # 下载失败，返回原路径
            except Exception as e:
                print(f"Failed to download image from {path}: {e}")
                return path

        # 本地路径处理
        if not os.path.isabs(path):
            potential_paths = [
                os.path.join(os.getcwd(), path),
                os.path.join(os.getcwd(), "web", path),
                os.path.join(os.getcwd(), "web", "images", os.path.basename(path)),
                os.path.join(os.getcwd(), "web", "images", "ai_generated", os.path.basename(path)),
                os.path.join(os.getcwd(), "web", "images", "web_crawled", os.path.basename(path))
            ]
            for p_path in potential_paths:
                if os.path.exists(p_path):
                    return p_path
        
        return path

    @staticmethod
    def _generate_docx(file_path: str, content: List[Dict[str, Any]], style_config: Optional[Dict[str, Any]]) -> str:
        doc = Document()
        
        # 如果可能，应用全局样式（python-docx 在没有模板的情况下全局默认设置有限）
        # 目前我们将逐个元素应用。
        
        for block in content:
            b_type = block.get("type", "paragraph")
            text = block.get("text", "")
            
            if b_type == "heading":
                level = block.get("level", 1)
                doc.add_heading(text, level=level)
                
            elif b_type == "paragraph":
                p = doc.add_paragraph()
                run = p.add_run(text)
                
                # 样式覆盖
                style = block.get("style", {})
                if "font_size" in style:
                    run.font.size = Pt(style["font_size"])
                if "bold" in style:
                    run.bold = style["bold"]
                if "italic" in style:
                    run.italic = style["italic"]
                if "alignment" in style:
                    align = style["alignment"].lower()
                    if align == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif align == "justify":
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
            elif b_type == "image":
                path = block.get("path")
                width = block.get("width", 6.0) # 默认 6 英寸
                
                if path:
                    path = DocumentGenerator._resolve_image_path(path)

                if path and os.path.exists(path):
                    try:
                        doc.add_picture(path, width=Inches(min(width, 6.0))) # 限制最大宽度
                    except Exception as e:
                        doc.add_paragraph(f"[Image load failed: {str(e)}]")
                else:
                    doc.add_paragraph(f"[Image not found: {block.get('path')}]")
                    
            elif b_type == "page_break":
                doc.add_page_break()
                
        doc.save(file_path)
        return f"Document generated successfully. Access path: /static/files/{os.path.basename(file_path)}"

    @staticmethod
    def _generate_pdf(file_path: str, content: List[Dict[str, Any]], style_config: Optional[Dict[str, Any]], downloaded_files: Optional[List[str]] = None) -> str:
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # 注册并应用中文字体
        chinese_font = DocumentGenerator._register_chinese_font()
        
        # 更新所有样式使用中文字体
        for style_name in styles.byName:
            styles[style_name].fontName = chinese_font
            # 确保字号合适
            if style_name.startswith('Heading'):
                styles[style_name].leading = styles[style_name].fontSize * 1.5
        
        # 可以根据 style_config 在此处添加自定义样式
        
        for block in content:
            b_type = block.get("type", "paragraph")
            text = block.get("text", "")
            
            if b_type == "heading":
                level = block.get("level", 1)
                # 将级别映射到 h1, h2, h3...
                style_name = f'Heading{min(level, 6)}'
                if style_name not in styles:
                    style_name = 'Heading1'
                story.append(Paragraph(text, styles[style_name]))
                
            elif b_type == "paragraph":
                # 如果需要特定样式，为此段落创建自定义样式
                # 为简单起见，使用 'BodyText' 或创建临时样式
                style = styles['BodyText']
                
                # ReportLab 通常通过文本中的类似 XML 的标签处理样式，
                # 或者我们可以定义一个新的 ParagraphStyle。
                # 对于基本文本，我们只使用 BodyText。
                # 如果需要对齐：
                block_style = block.get("style", {})
                if "alignment" in block_style:
                    # 克隆并修改
                    align_map = {"left": 0, "center": 1, "right": 2, "justify": 4}
                    align = block_style["alignment"].lower()
                    if align in align_map:
                        style = ParagraphStyle(
                            f'Custom_{len(story)}',
                            parent=styles['BodyText'],
                            alignment=align_map[align]
                        )
                
                story.append(Paragraph(text, style))
                story.append(Spacer(1, 12))
                
            elif b_type == "image":
                path = block.get("path")
                width = block.get("width", 400) # 大致像素
                height = block.get("height", 300)
                
                if path:
                    path = DocumentGenerator._resolve_image_path(path, downloaded_files)
                
                if path and os.path.exists(path):
                    try:
                        # ReportLab 图像
                        # 转换宽度概念（英寸/像素）？
                        # ReportLab 使用点（1/72 英寸）。
                        # 如果用户为 docx 传递 6（英寸），我们可能需要转换。
                        # 让我们假设输入是通用的 'width' 比例。
                        # 如果 < 10，假设为英寸 -> * 72。
                        img_width = width
                        if img_width < 20: 
                            img_width = img_width * 72
                            
                        # 保持纵横比的逻辑可以在这里实现
                        img = RLImage(path, width=img_width, height=img_width*0.75) # 纵横比猜测
                        # 更好：如果我们只设置宽度，让 reportlab 处理纵横比？
                        # RLImage(path, width=w, height=h) 需要两者或计算逻辑。
                        # 让我们尝试使用 PIL 获取图像大小以保持纵横比
                        try:
                            from PIL import Image as PILImage
                            with PILImage.open(path) as pi:
                                iw, ih = pi.size
                                aspect = ih / iw
                                img_height = img_width * aspect
                                img = RLImage(path, width=img_width, height=img_height)
                        except:
                            pass # 回退到近似方形
                            
                        story.append(img)
                        story.append(Spacer(1, 12))
                    except Exception as e:
                        print(f"Image load failed for {path}: {e}")
                        story.append(Paragraph(f"[Image load failed: {str(e)}]", styles['BodyText']))
                else:
                    print(f"Image not found during generation: {path}")
                    story.append(Paragraph(f"[Image not found: {block.get('path')}]", styles['BodyText']))
                    
            elif b_type == "page_break":
                from reportlab.platypus import PageBreak
                story.append(PageBreak())
                
        doc.build(story)
        return f"Document generated successfully. Access path: /static/files/{os.path.basename(file_path)}"
